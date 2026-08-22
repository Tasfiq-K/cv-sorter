from __future__ import annotations

from pathlib import Path

import pdfplumber
from docx import Document
from . models import CandidateDocument


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}


class UnsupportedFileTypeError(Exception):
    """Raised when attempting to extract text from an unsupported file type."""


def extract_text(file_path: str | Path) -> str:
    """
    Extract text from a supported document.

    Supported formats:
        - PDF
        - DOCX
        - TXT

    Parameters
    ----------
    file_path : str | Path
        Path to the document.

    Returns
    -------
    str
        Extracted text.

    Raises
    ------
    FileNotFoundError
        If the file does not exist.

    UnsupportedFileTypeError
        If the file extension is unsupported.
    """

    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")

    extension = path.suffix.lower()

    if extension == ".pdf":
        return _extract_pdf(path)

    if extension == ".docx":
        return _extract_docx(path)

    if extension == ".txt":
        return _extract_txt(path)

    raise UnsupportedFileTypeError(
        f"Unsupported file type '{extension}'. "
        f"Supported types: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
    )


def extract_directory(folder: str | Path) -> dict[str, str]:
    """
    Extract text from every supported document inside a directory.

    Returns
    -------
    dict[str, str]

        {
            "john.pdf": "...text...",
            "alice.docx": "...text..."
        }
    """

    folder = Path(folder)

    if not folder.exists():
        raise FileNotFoundError(folder)

    documents: list[CandidateDocument] = []

    for file in sorted(folder.iterdir()):
        if file.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue

        try:
            documents.append(extract_text(file))
            # documents[file.name] = extract_text(file)

        except Exception as e:
            print(f"Skipping {file.name}: {e}")

    return documents


# ---------------------------------------------------------------------
# Internal extractors
# ---------------------------------------------------------------------


def _extract_pdf(path: Path) -> str:
    pages: list[str] = []

    with pdfplumber.open(path) as pdf:
        page_count = len(pdf.pages)

        for page in pdf.pages:
            text = page.extract_text(
                x_tolerance=2,
                y_tolerance=3
            )

            if text:
                pages.append(text)

    raw_text = "\n\n".join(pages).strip()

    return CandidateDocument(
        filename=path.name,
        path=path,
        extension=path.suffix.lower(),
        raw_text=raw_text,
        page_count=page_count,
        character_count=len(raw_text),
    )

def _extract_docx(path: Path) -> str:
    doc = Document(path)

    paragraphs = [
        p.text.strip()
        for p in doc.paragraphs
        if p.text.strip()
    ]

    text = "\n".join(paragraphs)

    return CandidateDocument(
        filename=path.name,
        path=path,
        extension=path.suffix.lower(),
        raw_text=text,
        page_count=1,
        character_count=len(text)
    )



def _extract_txt(path: Path) -> str:
    text = path.read_text(...)

    return CandidateDocument(
        filename=path.name,
        path=path,
        extension=path.suffix.lower(),
        raw_text=text,
        page_count=1,
        character_count=len(text)
    )

# def _test_pdf_tolerance(path: Path) -> None:
#     with pdfplumber.open(path) as pdf:
#         for tolerance in [0.5, 1, 1.5, 2, 3, 4, 5]:

#             print("\n" + "=" * 80)
#             print(f"x_tolerance = {tolerance}")
#             print("=" * 80)

#             text = pdf.pages[0].extract_text(
#                 x_tolerance=tolerance,
#                 y_tolerance=3,
#             )

#             print(text[:3000] if text else "[NO TEXT]")


if __name__ == "__main__":
    from pprint import pprint
    file_path = "data/cvs/Md Tasfiq Kamran.pdf"
    text = extract_text(file_path)
    # file_path = Path("data/cvs/Md Tasfiq Kamran.pdf")
    # _test_pdf_tolerance(file_path)

    pprint(text)
